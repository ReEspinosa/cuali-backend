import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUTPUT_DIR = "generated_docs"

LETRAS = ["a", "b", "c", "d", "e", "f"]


def generar_docx_cuestionario(recurso_id: str, datos: dict) -> str:
    """
    datos = {
      "titulo": str,
      "preguntas": [
        { "pregunta": str, "tipo": "opcion_multiple|verdadero_falso|abierta",
          "opciones": [str] | None, "respuesta_correcta": str | None,
          "respuesta_sugerida": str | None }
      ]
    }
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    path = os.path.join(OUTPUT_DIR, f"cuestionario_{recurso_id}.docx")

    doc = Document()

    titulo_p = doc.add_paragraph()
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo_p.add_run(datos.get("titulo", "Cuestionario"))
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x3F, 0x5B, 0x45)

    doc.add_paragraph()
    nombre_p = doc.add_paragraph()
    nombre_p.add_run("Nombre: _____________________________________     Fecha: ____________").font.size = Pt(11)
    doc.add_paragraph()

    for i, p in enumerate(datos.get("preguntas", []), start=1):
        pregunta_p = doc.add_paragraph()
        run = pregunta_p.add_run(f"{i}. {p.get('pregunta', '')}")
        run.bold = True
        run.font.size = Pt(12)

        tipo = p.get("tipo", "abierta")
        if tipo in ("opcion_multiple", "verdadero_falso"):
            for j, opcion in enumerate(p.get("opciones", [])):
                letra = LETRAS[j] if j < len(LETRAS) else str(j + 1)
                op_p = doc.add_paragraph(style="List Bullet")
                op_p.paragraph_format.left_indent = Pt(24)
                op_p.add_run(f"{letra}) {opcion}").font.size = Pt(11)
        else:
            for _ in range(3):
                linea = doc.add_paragraph()
                linea.add_run("_______________________________________________________").font.size = Pt(11)

        doc.add_paragraph()

    doc.add_page_break()
    hoja_p = doc.add_paragraph()
    hoja_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hoja_p.add_run("Hoja de respuestas (solo para el maestro)")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x3F, 0x5B, 0x45)
    doc.add_paragraph()

    for i, p in enumerate(datos.get("preguntas", []), start=1):
        tipo = p.get("tipo", "abierta")
        resp_p = doc.add_paragraph()
        if tipo == "abierta":
            resp_p.add_run(f"{i}. Respuesta sugerida: ").bold = True
            resp_p.add_run(p.get("respuesta_sugerida", "(respuesta libre)"))
        else:
            resp_p.add_run(f"{i}. Respuesta correcta: ").bold = True
            resp_p.add_run(p.get("respuesta_correcta", ""))

    doc.save(path)
    return path