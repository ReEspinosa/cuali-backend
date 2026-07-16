"""
Extrae texto de archivos adjuntos para poder pasárselo al LLM como contexto.

Soporta:
- .docx  -> texto de párrafos y tablas (via python-docx)
- .pdf   -> texto por página (via pypdf)
- .png/.jpg/.jpeg -> el modelo actual (gpt-oss-20b vía LM Studio) NO tiene
  entrada de visión, así que no podemos "leer" imágenes todavía. Se
  regresa None y el llamador debe avisarle al maestro que no se procesó.
- .doc (Word viejo, binario) -> tampoco soportado directamente por
  python-docx; se regresa None con el mismo aviso.

Cuando conectes un modelo con visión (o un servicio de OCR), agrega el caso
de imagen aquí sin tocar el resto del flujo.
"""

from pathlib import Path

LIMITE_CARACTERES = 12000  # evita mandar documentos gigantes al LLM


def _extraer_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    partes = [p.text for p in doc.paragraphs if p.text.strip()]

    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells]
            if any(celdas):
                partes.append(" | ".join(celdas))

    return "\n".join(partes)


def _extraer_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    partes = [pagina.extract_text() or "" for pagina in reader.pages]
    return "\n".join(partes)


def extraer_texto(path: Path) -> str | None:
    extension = path.suffix.lower()

    try:
        if extension == ".docx":
            texto = _extraer_docx(path)
        elif extension == ".pdf":
            texto = _extraer_pdf(path)
        else:
            # .doc, .png, .jpg, .jpeg u otros: no soportado por ahora.
            print(f"[extraccion_archivos] Extensión no soportada para extracción: {extension} ({path})")
            return None
    except Exception as exc:
        # TEMPORAL: imprime el error real en la consola de uvicorn para
        # diagnosticar. Cuando confirmemos que funciona bien, esto se puede
        # quitar o mandar a un logger real en vez de print().
        print(f"[extraccion_archivos] ERROR extrayendo texto de {path} ({extension}): {exc!r}")
        return None

    texto = texto.strip()
    if not texto:
        print(f"[extraccion_archivos] Extracción vacía para {path} ({extension}) — el archivo no tenía texto detectable.")
        return None

    print(f"[extraccion_archivos] Extraídos {len(texto)} caracteres de {path} ({extension}).")

    if len(texto) > LIMITE_CARACTERES:
        texto = texto[:LIMITE_CARACTERES] + "\n[...documento truncado por longitud...]"

    return texto