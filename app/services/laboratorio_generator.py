import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUTPUT_DIR = "generated_docs"


def generar_docx_laboratorio(recurso_id: str, titulo: str, contenido: str) -> str:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    path = os.path.join(OUTPUT_DIR, f"laboratorio_{recurso_id}.docx")

    doc = Document()

    titulo_p = doc.add_paragraph()
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo_p.add_run(titulo)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x3F, 0x5B, 0x45)

    doc.add_paragraph()

    for linea in contenido.split("\n"):
        linea = linea.strip()
        if not linea:
            doc.add_paragraph()
            continue
        if linea.startswith("- "):
            doc.add_paragraph(linea[2:], style="List Bullet")
        else:
            doc.add_paragraph(linea)

    doc.save(path)
    return path