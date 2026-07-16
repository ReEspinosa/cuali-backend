import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUTPUT_DIR = "generated_docs"


def _shade_cell(cell, color_hex: str) -> None:
    shading = cell._tc.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex})
    shading.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _info_row(table, label: str, value: str) -> None:
    row = table.add_row()
    _set_cell_text(row.cells[0], label, bold=True)
    _shade_cell(row.cells[0], "EDEAE3")
    _set_cell_text(row.cells[1], value)


def generar_docx_planeacion(planeacion_id: str, datos: dict) -> str:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    path = os.path.join(OUTPUT_DIR, f"planeacion_{planeacion_id}.docx")

    doc = Document()

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("Avance - Plan de Clase")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x3F, 0x5B, 0x45)

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.add_run(
        f"{datos['grado']}. \"{datos['grupo']}\"  |  {datos['campo_formativo']}"
    ).font.size = Pt(11)

    doc.add_paragraph()

    info = doc.add_table(rows=0, cols=2)
    info.style = "Table Grid"
    info.columns[0].width = Pt(140)
    info.columns[1].width = Pt(360)

    _info_row(info, "Campo Formativo", datos["campo_formativo"])
    _info_row(info, "Metodologia", datos["metodologia"])
    _info_row(info, "Grado / Grupo", f"{datos['grado']}. \"{datos['grupo']}\"")
    _info_row(info, "Ejes articuladores", "\n".join(datos["ejes_articuladores"]))
    _info_row(info, "Titulo del proyecto", datos["titulo_proyecto"])
    _info_row(info, "Contenido", datos["contenido"])
    _info_row(info, "Proceso de desarrollo de aprendizaje", datos["pda"])
    _info_row(info, "Intencion didactica", datos["intencion_didactica"])
    _info_row(info, "Proposito", datos["proposito"])

    doc.add_paragraph()

    dias_table = doc.add_table(rows=1, cols=5)
    dias_table.style = "Table Grid"
    headers = ["Etapa", "Sesion", "Actividades", "Recursos", "Tarea"]
    for i, h in enumerate(headers):
        cell = dias_table.rows[0].cells[i]
        _set_cell_text(cell, h, bold=True, size=10)
        _shade_cell(cell, "3F5B45")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    etapa_anterior = None
    etapa_cell_inicio = None
    for fila in datos["dias"]:
        row = dias_table.add_row()
        _set_cell_text(row.cells[0], fila.get("etapa", ""), bold=True, size=9)
        _set_cell_text(row.cells[1], fila.get("sesion", ""), bold=True, size=9)
        _set_cell_text(row.cells[2], fila.get("actividades", ""), size=9)
        _set_cell_text(row.cells[3], fila.get("recursos", ""), size=9)
        _set_cell_text(row.cells[4], fila.get("tarea", ""), size=9)

        etapa_actual = fila.get("etapa", "")
        if etapa_actual and etapa_actual == etapa_anterior:
            etapa_cell_inicio.merge(row.cells[0])
        else:
            etapa_cell_inicio = row.cells[0]
        etapa_anterior = etapa_actual

    doc.add_paragraph()

    eval_table = doc.add_table(rows=1, cols=2)
    eval_table.style = "Table Grid"
    _set_cell_text(eval_table.rows[0].cells[0], "Evaluacion", bold=True)
    _shade_cell(eval_table.rows[0].cells[0], "EDEAE3")
    _set_cell_text(eval_table.rows[0].cells[1], datos["evaluacion"])

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("___________________________")
    firma = doc.add_paragraph()
    firma.add_run("Maestra de grupo").bold = True

    doc.save(path)
    return path
